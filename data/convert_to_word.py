"""
Convert README.md to Word document (ProjectDetails.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Add a heading with formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph(text)
    if bold:
        para.runs[0].bold = True
    if italic:
        para.runs[0].italic = True
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    doc.add_paragraph(text, style='List Bullet')

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph(text)
    para.runs[0].font.name = 'Courier New'
    para.runs[0].font.size = Pt(9)
    return para

def convert_readme_to_word():
    """Convert README.md to Word document"""

    # Read README.md
    with open('README.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set document title
    title = doc.add_heading('AI Project: Sentiment Analysis and Review Classification System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph('Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True

    doc.add_paragraph()  # Empty line

    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        # Skip the first title line (already added)
        if line.startswith('# AI Project:'):
            continue

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    add_code_block(doc, '\n'.join(code_lines))
                    code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle headings
        if line.startswith('## '):
            add_heading(doc, line[3:], level=1)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=2)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=3)

        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)

        # Handle bullet points
        elif line.startswith('- '):
            add_bullet_point(doc, line[2:])

        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line, style='List Number')

        # Handle bold text with **
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are bold
                    run = para.add_run(part)
                    run.bold = True
                else:
                    para.add_run(part)

        # Handle regular paragraphs
        elif line.strip():
            doc.add_paragraph(line)

    # Add page breaks between major sections for better formatting

    # Save document
    output_path = 'ProjectDetails.docx'
    doc.save(output_path)
    print(f"✅ Word document created successfully: {output_path}")
    print(f"📄 Location: /Users/sahilijaz/Desktop/ai-project/{output_path}")

if __name__ == "__main__":
    convert_readme_to_word()
